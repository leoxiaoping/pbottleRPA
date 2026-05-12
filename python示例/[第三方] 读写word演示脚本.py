"""
PBottle RPA demo. Please refer to the *process documentation* for specific APIs.
Official website: https://rpa.pbottle.com/
Process documentation: https://rpa.pbottle.com/docs/
"""

import sys
import os
import time

# Add parent directory to path to import pbottleRPA
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import pbottleRPA as rpa

# Try to import python‑docx and mammoth
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import mammoth
except ImportError:
    rpa.showMsg(
        "Please install third‑party modules", "Run: pip install python-docx mammoth"
    )
    rpa.tts("Please install third‑party modules. Run: pip install python-docx mammoth")
    rpa.exit("Please install third‑party modules. Run: pip install python-docx mammoth")


print("=== Backend Word Read/Write Test ===")
print(rpa.getTime())
rpa.tts("Backend Word Read/Write Test")
rpa.wait(3)
rpa.tts("Generating a Word document in the background")
rpa.wait(5)

# Generate a Word document
doc = Document()

# Add a title paragraph
heading = doc.add_heading("Title Text", level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = heading.runs[0]
run.bold = True
run.font.size = Pt(20)

# Add a normal paragraph
p = doc.add_paragraph()
run1 = p.add_run("PBottle RPA official website: ")
run1.font.size = Pt(12)

# Add a hyperlink‑style text
hyperlink_run = p.add_run("https://rpa.pbottle.com")
hyperlink_run.font.size = Pt(12)
hyperlink_run.font.underline = True

# Add more sample content
doc.add_paragraph()
doc.add_paragraph("This is a Python‑based Word read/write demonstration.")
doc.add_paragraph("PBottle RPA supports a variety of automation operations, including:")
doc.add_paragraph("• Mouse and keyboard operations", style="List Bullet")
doc.add_paragraph("• Image recognition", style="List Bullet")
doc.add_paragraph("• AI capability integration", style="List Bullet")
doc.add_paragraph("• File operations", style="List Bullet")

# Save the file
word_path = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "Word_test_document.docx",
)
doc.save(word_path)
rpa.openDir(os.path.dirname(word_path))
rpa.tts("Word_test_document.docx has been generated. Please check.")
rpa.wait(3)

# Read the Word document
rpa.tts("Reading the Word document in the background and displaying content in the log")
rpa.wait(3)

# Use mammoth to extract text content
with open(word_path, "rb") as f:
    result = mammoth.extract_raw_text(f)
    print("Word document content:", result.value)

rpa.tts("Word test document has been read and logged.")
rpa.wait(2)


if __name__ == "__main__":
    pass
